import spacy
import re
import docx
import streamlit as st

# load spaCy NLP model
nlp = spacy.load("en_core_web_sm")

# regex patterns
contact_pattern = r"\b(?:\d[ -]*?){8,}\b"
date_pattern = r"\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b"
email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"

def replace_personal_info(text):
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            text = text.replace(ent.text, "NAME")
    text = re.sub(contact_pattern, "CONTACT NUMBER", text)
    text = re.sub(date_pattern, "DATE", text)
    text = re.sub(email_pattern, "EMAIL ADDRESS", text)
    return text

# Streamlit app
def app():
    # file uploader
    st.header("Upload a Word document to remove personal information:")
    uploaded_file = st.file_uploader("Choose a Word document", type=["docx"])

    # remove personal information button
    if uploaded_file is not None:
        st.header("Original Document:")
        doc = docx.Document(uploaded_file)
        for para in doc.paragraphs:
            st.write(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    st.write(cell.text)

        if st.button("Remove Personal Information"):
            # create a new document to save the modified content
            new_doc = docx.Document()
            for para in doc.paragraphs:
                new_doc.add_paragraph(replace_personal_info(para.text))
            for table in doc.tables:
                new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.cell(i, j).text = replace_personal_info(cell.text)

            # download the modified document
            st.header("Modified Document:")
            modified_doc_name = "modified_" + uploaded_file.name
            new_doc.save(modified_doc_name)
            with open(modified_doc_name, "rb") as f:
                bytes_data = f.read()
            st.download_button(label="Download Modified Document", data=bytes_data, file_name=modified_doc_name)

if __name__ == "__main__":
    app()
